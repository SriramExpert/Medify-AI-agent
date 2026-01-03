import os
import tempfile
from typing import Dict, List, Any, Optional
import logging
from pathlib import Path

# Document processing libraries
import PyPDF2
import docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_classic.chains import RetrievalQA
from langchain_classic.llms import OpenAI
from langchain_core.documents import Document as LangchainDocument
import numpy as np

# Web search fallback
from googlesearch import search as google_search
import requests
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)

class DocumentTool:
    """Tool for document processing and web search"""
    
    def __init__(self, openai_api_key: Optional[str] = None):
        self.openai_api_key = openai_api_key
        self.vector_store = None
        self.documents = []
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {e}")
            raise
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise
    
    def process_uploaded_file(self, file_path: str, file_extension: str) -> str:
        """Process uploaded file based on extension"""
        if file_extension.lower() == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension.lower() == '.txt':
            return self.extract_text_from_txt(file_path)
        elif file_extension.lower() == '.docx':
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def create_vector_store(self, text: str, document_name: str = "uploaded_document"):
        """Create vector store from document text"""
        try:
            # Split text into chunks
            texts = self.text_splitter.split_text(text)
            
            # Create documents
            docs = [LangchainDocument(page_content=t, metadata={"source": document_name}) for t in texts]
            self.documents = docs
            
            # Create embeddings and vector store
            if self.openai_api_key:
                embeddings = OpenAIEmbeddings(openai_api_key=self.openai_api_key)
                self.vector_store = FAISS.from_documents(docs, embeddings)
            else:
                # Fallback: simple similarity search without embeddings
                logger.warning("OpenAI API key not provided, using simple text search")
                self.vector_store = SimpleVectorStore(docs)
            
            logger.info(f"Created vector store with {len(docs)} chunks")
            return True
        except Exception as e:
            logger.error(f"Error creating vector store: {e}")
            return False
    
    def query_document(self, question: str, top_k: int = 3) -> Dict[str, Any]:
        """Query the document with a question"""
        try:
            if not self.vector_store or not self.documents:
                return {
                    "success": False,
                    "error": "No document loaded",
                    "answer": None,
                    "source": None
                }
            
            if self.openai_api_key:
                # Use OpenAI for better answers
                llm = OpenAI(openai_api_key=self.openai_api_key, temperature=0)
                qa_chain = RetrievalQA.from_chain_type(
                    llm=llm,
                    chain_type="stuff",
                    retriever=self.vector_store.as_retriever(search_kwargs={"k": top_k})
                )
                result = qa_chain({"query": question})
                answer = result["result"]
                source = "AI-enhanced document search"
            else:
                # Simple similarity search
                answer, source = self._simple_similarity_search(question, top_k)
            
            # Check if answer is confident
            confidence = self._calculate_confidence(answer, question)
            
            return {
                "success": True,
                "answer": answer,
                "source": source,
                "confidence": confidence,
                "from_document": confidence > 0.3  # Threshold for document relevance
            }
        except Exception as e:
            logger.error(f"Error querying document: {e}")
            return {
                "success": False,
                "error": str(e),
                "answer": None,
                "source": None,
                "from_document": False
            }
    
    def _simple_similarity_search(self, question: str, top_k: int = 3) -> tuple:
        """Simple similarity search without embeddings"""
        question_lower = question.lower()
        question_words = set(question_lower.split())
        
        scores = []
        for doc in self.documents:
            content_lower = doc.page_content.lower()
            content_words = set(content_lower.split())
            
            # Simple word overlap score
            overlap = len(question_words.intersection(content_words))
            total_unique = len(question_words.union(content_words))
            score = overlap / total_unique if total_unique > 0 else 0
            
            scores.append((score, doc.page_content, doc.metadata["source"]))
        
        # Sort by score and get top results
        scores.sort(reverse=True, key=lambda x: x[0])
        top_results = scores[:top_k]
        
        if top_results and top_results[0][0] > 0:
            # Combine top results
            combined_answer = "\n".join([f"- {content[:200]}..." for score, content, source in top_results if score > 0])
            source = f"From document: {top_results[0][2]}"
            return combined_answer, source
        else:
            return "No relevant information found in the document.", "Document search"
    
    def _calculate_confidence(self, answer: str, question: str) -> float:
        """Calculate confidence score for the answer"""
        if not answer or "not found" in answer.lower() or "no relevant" in answer.lower():
            return 0.0
        
        # Simple confidence calculation based on answer length and question keywords
        question_words = set(question.lower().split())
        answer_words = set(answer.lower().split())
        
        overlap = len(question_words.intersection(answer_words))
        confidence = min(overlap / len(question_words) if question_words else 0, 1.0)
        
        # Boost confidence for longer answers that seem substantive
        if len(answer.split()) > 10:
            confidence = min(confidence + 0.2, 1.0)
        
        return confidence
    
    def web_search(self, query: str, num_results: int = 3) -> List[Dict[str, str]]:
        """Perform web search as fallback"""
        try:
            results = []
            search_results = list(google_search(query, num_results=num_results, advanced=True))
            
            for i, result in enumerate(search_results[:num_results]):
                try:
                    # Fetch page content
                    response = requests.get(result.url, timeout=5)
                    soup = BeautifulSoup(response.content, 'html.parser')
                    
                    # Extract relevant text (first 500 chars)
                    text = soup.get_text()[:500] + "..."
                    
                    results.append({
                        "title": result.title,
                        "url": result.url,
                        "snippet": result.description,
                        "content": text,
                        "rank": i + 1
                    })
                except:
                    # If can't fetch content, use description
                    results.append({
                        "title": result.title,
                        "url": result.url,
                        "snippet": result.description,
                        "content": result.description,
                        "rank": i + 1
                    })
            
            return results
        except Exception as e:
            logger.error(f"Error in web search: {e}")
            # Return mock results for demonstration
            return self._get_mock_search_results(query)
    
    def _get_mock_search_results(self, query: str) -> List[Dict[str, str]]:
        """Return mock search results for demonstration"""
        return [
            {
                "title": f"Information about {query}",
                "url": "https://example.com/search",
                "snippet": f"Search results for {query}. This is a mock response.",
                "content": f"This is a mock response for the query: {query}. In a real implementation, this would be actual web search results.",
                "rank": 1
            }
        ]

class SimpleVectorStore:
    """Simple in-memory vector store for fallback"""
    def __init__(self, documents):
        self.documents = documents
    
    def as_retriever(self, search_kwargs=None):
        return self
    
    def get_relevant_documents(self, query):
        # Simple keyword matching
        query_lower = query.lower()
        relevant = []
        for doc in self.documents:
            if any(word in doc.page_content.lower() for word in query_lower.split()):
                relevant.append(doc)
        return relevant[:search_kwargs.get('k', 3)] if search_kwargs else relevant[:3]